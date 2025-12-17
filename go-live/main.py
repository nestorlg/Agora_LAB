import docx
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm

import openpyxl as op
import configparser as cp

import Global as glb


def get_or_create_hyperlink_style(d):
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


with open("temp.txt", "r") as f:
    lineas = f.readlines()

project = lineas[0][:-1]

glb.datos_import = "projects/" + project + "/entrada.xlsx"
glb.datos_export = "projects/" + project + "/salida.xlsx"
glb.dir_imagenes = "imagenesAgora"

glb.ruta_cache = "outPutExcelCreator/data/cache.xlsx"

doc = Document()

"""
LECTURA DE DATOS DESDE EL ARCHIVO TXT
"""

archivo_txt = open("go-live/products-and-prices.txt", "r")

datos_directos = []

# Lectura de cada linea del fichero txt
for line in archivo_txt:
    fila = []
    for dato in line.split("|"):
        fila.append(dato.strip())

    if fila == ['']:
        break

    datos_directos.append(fila)

# Obtención de las distintas tarifas contempladas
tarifas = []

for dato in datos_directos:
    if dato[3] not in tarifas:
        tarifas.append(dato[3])
    else:
        break

nombres_columnas = [
    'Producto',
    'Familia',
    'Impuesto',
    'P. Princ.',
    'P. Añadido'
]

n_tarifas = len(tarifas)
n_filas = int(len(datos_directos) // n_tarifas) + 1
n_columnas = len(nombres_columnas)

doc.alignmet = WD_ALIGN_PARAGRAPH.CENTER
run = doc.add_picture("go-live/logo_infogral.jpg")
p = doc.add_paragraph()

p_format = p.paragraph_format
p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

entrada = op.load_workbook(glb.datos_import)
hoja_datos = entrada["Datos adicionales"]

"""
0- Nombre comercial
1- Nombre fiscal
2- CIF
3- Calle
4- Localidad
5- Ciudad
6- Código postal
"""

datos_fiscales = [
    hoja_datos["B2"].value if hoja_datos["B2"].value is not None else "",
    hoja_datos["B1"].value,
    hoja_datos["B3"].value,
    hoja_datos["B4"].value if hoja_datos["B4"].value is not None else "",
    hoja_datos["B6"].value if hoja_datos["B6"].value is not None else "",
    hoja_datos["B7"].value if hoja_datos["B7"].value is not None else "",
    str(hoja_datos["B5"].value) if hoja_datos["B5"].value is not None else "",
]

run = p.add_run(datos_fiscales[0] + "\n")
run = p.add_run(datos_fiscales[1] + " | " + datos_fiscales[2] + "\n")
run = p.add_run(datos_fiscales[3] + ", " + datos_fiscales[6] + ", " + datos_fiscales[4] + ", " + datos_fiscales[5])
font = run.font
font.size = Pt(12)

doc.add_page_break()

p_elm = p._element
for border_name in ('top', 'left', 'bottom', 'right'):
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '12')
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), '000000')
    p_elm.get_or_add_pPr().append(border)

for tarifa in tarifas:
    doc.add_heading("Catálogo - " + tarifa)
    tabla = doc.add_table(rows=n_filas, cols=n_columnas, style='Table Grid')

    for c in range(n_columnas):
        tabla.cell(0, c).text = nombres_columnas[c]
        tabla.cell(0, c).paragraphs[0].runs[0].font_size = Pt(12)
        tabla.cell(0, c).paragraphs[0].paragraph_format.alignment = 1

        for paragraph in tabla.cell(0, c).paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(128, 32, 4)

    for i in range(len(datos_directos)):
        if datos_directos[i][3] == tarifa:
            for j in range(len(datos_directos[i])):
                if j < 3:
                    this_col = j
                    tabla.cell(i // n_tarifas + 1, j).text = datos_directos[i][j]
                elif j > 3:
                    this_col = j - 1
                    if datos_directos[i][j] != "NULL":
                        tabla.cell(i // n_tarifas + 1, j - 1).text =\
                            str("{:.2f}".format(round(float(datos_directos[i][j]), 2)))
                    else:
                        tabla.cell(i // n_tarifas + 1, j - 1).text = "-"
                else:
                    continue

                tabla.cell(i // n_tarifas + 1, this_col).paragraphs[0].runs[0].font_size = Pt(10)
                tabla.cell(i // n_tarifas + 1, this_col).paragraphs[0].paragraph_format.alignment = 0

                for paragraph in tabla.cell(i // n_tarifas + 1, this_col).paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.bold = False

doc.add_heading("Manuales")

links = [
    ("Manual de usuario de Ágora", "https://www.agorapos.com/manual/agora-restaurant/manual-agora-restaurant.pdf"),
    ("Agora Training Channel", "https://www.youtube.com/@agora-softwaretpv-training3659")
]

for text, url in links:
    p = doc.add_paragraph(style="ListBullet")
    add_hyperlink(p, text, url)

# Add Infogral icon and page number

for i, section in enumerate(doc.sections):
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("go-live/icono_infogral.jpg", width=Inches(1.0))
    run.add_text("\n")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Página "
    run = paragraph.add_run()

    # Crear el campo de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

doc.save("projects/" + project + "/go-live.docx")
